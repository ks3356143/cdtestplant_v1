from pathlib import Path
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from utils.path_utils import project_path
from ninja import File, UploadedFile
from ninja.errors import HttpError
from ninja_extra.controllers import api_controller, ControllerBase, route
from ninja_jwt.authentication import JWTAuth
from ninja_extra.permissions import IsAuthenticated
from django.db import transaction
from django.shortcuts import get_object_or_404
from django.db.models import QuerySet
from docx import Document
from docxtpl import DocxTemplate
# 工具
from apps.createSeiTaiDocument.docXmlUtils import generate_temp_doc, get_frag_from_document
from apps.createSeiTaiDocument.schema import SeitaiInputSchema
from utils.chen_response import ChenResponse
from apps.project.models import Project, Dut
from apps.createDocument.extensions.documentTime import DocTime
from utils.util import get_str_dict
from apps.createSeiTaiDocument.extensions.download_response import get_file_respone
# 图片工具docx
from apps.createSeiTaiDocument.extensions.shape_size_tool import set_shape_size
# 修改temp文本片段工具
from apps.createSeiTaiDocument.docXmlUtils import get_jinja_stdContent_element, stdContent_modify

main_download_path = Path(settings.BASE_DIR) / 'media'

# @api_controller("/create", tags=['生成产品文档接口'], auth=JWTAuth(), permissions=[IsAuthenticated])
@api_controller("/create", tags=['生成产品文档接口'])
class GenerateSeitaiController(ControllerBase):
    chinese_round_name: list = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.project_obj: Project | None = None
        self.temp_context = {}

    @route.post("/dgDocument", url_name="create-dgDocument")
    @transaction.atomic
    def create_dgDocument(self, payload: SeitaiInputSchema):
        # 获取项目Model
        self.project_obj = get_object_or_404(Project, id=payload.id)
        # 生成大纲需要的文本片段信息储存在字典里面
        sec_title = get_str_dict(self.project_obj.secret, 'secret')
        duty_person = self.project_obj.duty_person
        is_jd = True if self.project_obj.report_type == '9' else False
        self.temp_context = {
                                'is_jd': is_jd,
                                'jd_or_third': "鉴定" if is_jd else "第三方",
                                'project_ident': self.project_obj.ident,
                                'project_name': self.project_obj.name,
                                'test_purpose': "装备鉴定和列装定型" if is_jd else "软件交付和使用",
                                'sec_title': '密级：' + sec_title,
                                'sec': sec_title,
                                'duty_person': duty_person,
                                'member': self.project_obj.member[0] if len(
                                    self.project_obj.member) > 0 else duty_person,
                                'entrust_unit': self.project_obj.entrust_unit
                            } | DocTime(payload.id).dg_final_time()  # python3.9以上推荐使用|运算符合并
        # 调用self添加temp_context信息
        self.get_first_round_code_ident()
        self.get_xq_doc_informations()
        result = generate_temp_doc('dg', payload.id, frag_list=payload.frag)
        if isinstance(result, dict):
            return ChenResponse(status=400, code=400,
                                message=result.get('msg', 'dg未报出错误原因，反正在生成文档出错'))
        dg_replace_path, dg_seitai_final_path = result
        # ~~~~start：2025/04/19-新增渲染单个字段（可能封装为函数-对temp文件下的jinja字段处理）~~~~
        # 现在已经把alias和stdContent对应起来了
        text_frag_name_list, doc_docx = get_jinja_stdContent_element(dg_replace_path)
        # 遍历找出来的文本片段进行修改
        self.text_frag_replace_handle(text_frag_name_list, doc_docx)
        # ~~~~end~~~~
        try:
            doc_docx.save(dg_seitai_final_path)
            # 文件下载
            return get_file_respone(payload.id, '测评大纲')
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="文档未生成或生成错误！，{0}".format(e))

    @route.post('/smDocument', url_name='create-smDocument')
    @transaction.atomic
    def create_smDocument(self, payload: SeitaiInputSchema):
        """生成最后说明文档"""
        # 获取项目对象
        self.project_obj = get_object_or_404(Project, id=payload.id)
        # 首先第二层模版所需变量
        is_jd = True if self.project_obj.report_type == '9' else False
        self.temp_context = {
                                'project_name': self.project_obj.name,
                                'project_ident': self.project_obj.ident,
                                'is_jd': is_jd,
                                'jd_or_third': "鉴定" if is_jd else "第三方",
                                'ident': self.project_obj.ident,
                                'sec_title': get_str_dict(self.project_obj.secret, 'secret'),
                                'sec': get_str_dict(self.project_obj.secret, 'secret'),
                                'duty_person': self.project_obj.duty_person,
                                'member': self.project_obj.member[0] if len(
                                    self.project_obj.member) > 0 else self.project_obj.duty_person,
                            } | DocTime(payload.id).sm_final_time()
        self.get_first_round_code_ident()
        # 文档片段操作
        result = generate_temp_doc('sm', payload.id, frag_list=payload.frag)
        if isinstance(result, dict):
            return ChenResponse(code=400, status=400, message=result.get('msg', '无错误原因'))
        sm_to_tpl_file, sm_seitai_final_file = result

        # 文本片段操作
        text_frag_name_list, doc_docx = get_jinja_stdContent_element(sm_to_tpl_file)
        self.text_frag_replace_handle(text_frag_name_list, doc_docx)
        # 注册时间变量
        try:
            doc_docx.save(sm_seitai_final_file)
            return get_file_respone(payload.id, '测试说明')
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

    @route.post('/jlDocument', url_name='create-jlDocument')
    @transaction.atomic
    def create_jlDocument(self, payload: SeitaiInputSchema):
        self.project_obj = get_object_or_404(Project, id=payload.id)
        # seitai文档所需变量
        is_jd = True if self.project_obj.report_type == '9' else False
        member = self.project_obj.member[0] if len(
            self.project_obj.member) > 0 else self.project_obj.duty_person
        self.temp_context = {
                                'project_name': self.project_obj.name,
                                'project_ident': self.project_obj.ident,
                                'is_jd': is_jd,
                                'name': self.project_obj.name,
                                'ident': self.project_obj.ident,
                                'sec_title': get_str_dict(self.project_obj.secret, 'secret'),
                                'duty_person': self.project_obj.duty_person, 'member': member
                            } | DocTime(payload.id).jl_final_time()
        self.get_xq_doc_informations()  # 添加文本片段“xq_version”
        result = generate_temp_doc('jl', payload.id, frag_list=payload.frag)
        if isinstance(result, dict):
            return ChenResponse(code=400, status=400, message=result.get('msg', '无错误原因'))
        jl_to_tpl_file, jl_seitai_final_file = result
        text_frag_name_list, doc_docx = get_jinja_stdContent_element(jl_to_tpl_file)
        # 文本片段操作
        self.text_frag_replace_handle(text_frag_name_list, doc_docx)
        # 重新设置时序图大小
        for shape in doc_docx.inline_shapes:
            set_shape_size(shape)
        try:
            doc_docx.save(jl_seitai_final_file)
            return get_file_respone(payload.id, '测试记录')
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

    @route.post('/hsmDocument', url_name='create-hsmDocument')
    @transaction.atomic
    def create_hsmDocument(self, payload: SeitaiInputSchema):
        """生成最后的回归测试说明-（多个文档）"""
        self.project_obj = get_object_or_404(Project, id=payload.id)
        hround_list: QuerySet = self.project_obj.pField.exclude(key='0')  # 非第一轮次
        cname_list = []
        if len(hround_list) < 1:
            return ChenResponse(code=400, status=400, message='无回归轮次，请添加后再生成')
        for hround in hround_list:
            # 获取当前轮次中文数字
            cname = self.chinese_round_name[int(hround.key)]
            # 将cname存入一个list，以便后续拼接给下载函数
            cname_list.append(cname)
            is_jd = True if self.project_obj.report_type == '9' else False
            member = self.project_obj.member[0] if len(
                self.project_obj.member) > 0 else self.project_obj.duty_person
            # 回归轮次的标识和版本
            so_dut: Dut = hround.rdField.filter(type='SO').first()
            if not so_dut:
                return ChenResponse(status=400, code=400, message=f'您缺少第{cname}轮的源代码被测件')
            # 每次循环会更新temp_context
            self.temp_context = {
                                    'project_name': self.project_obj.name,
                                    'project_ident': self.project_obj.ident,
                                    'is_jd': is_jd,
                                    'sec_title': get_str_dict(self.project_obj.secret, 'secret'),
                                    'duty_person': self.project_obj.duty_person,
                                    'member': member,
                                    'round_num': str(int(hround.key) + 1),
                                    'round_num_chn': cname,
                                    'soft_ident': so_dut.ref,
                                    'soft_version': so_dut.version,
                                    'location': hround.location,
                                } | DocTime(payload.id).hsm_final_time(hround.key)
            # 注意回归测试说明、回归测试记录都生成多个文档
            result = generate_temp_doc('hsm', payload.id, round_num=cname, frag_list=payload.frag)
            if isinstance(result, dict):
                return ChenResponse(status=400, code=400,
                                    message=result.get('msg', '回归测试说明生成报错...'))
            hsm_replace_path, hsm_seitai_final_path = result
            text_frag_name_list, doc_docx = get_jinja_stdContent_element(hsm_replace_path)
            # 文本片段操作
            self.text_frag_replace_handle(text_frag_name_list, doc_docx)
            try:
                doc_docx.save(hsm_seitai_final_path)
            except PermissionError as e:
                return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))
        # 因为回归说明、回归记录可能有多份，多份下载zip否则docx
        if len(cname_list) == 1:
            return get_file_respone(payload.id, '第二轮回归测试说明')
        else:
            return get_file_respone(payload.id, list(map(lambda x: f"第{x}轮回归测试说明", cname_list)))

    @route.post('/hjlDocument', url_name='create-hjlDocument')
    @transaction.atomic
    def create_hjlDocument(self, payload: SeitaiInputSchema):
        """生成最后的回归测试记录-（多个文档）"""
        self.project_obj: Project = get_object_or_404(Project, id=payload.id)
        # 取非第一轮次
        hround_list: QuerySet = self.project_obj.pField.exclude(key='0')
        cname_list = []
        if len(hround_list) < 1:
            return ChenResponse(code=400, status=400, message='无回归测试轮次，请创建后再试')
        for hround in hround_list:
            # 取出当前轮次key减1就是上一轮次
            cname = self.chinese_round_name[int(hround.key)]  # 输出二、三...
            cname_list.append(cname)
            member = self.project_obj.member[0] if len(
                self.project_obj.member) > 0 else self.project_obj.duty_person
            is_jd = True if self.project_obj.report_type == '9' else False
            so_dut: Dut = hround.rdField.filter(type='SO').first()
            if not so_dut:
                return ChenResponse(status=400, code=400, message=f'您缺少第{cname}轮的源代码被测件')
            self.temp_context = {
                                    'project_name': self.project_obj.name,
                                    'project_ident': self.project_obj.ident,
                                    'is_jd': is_jd,
                                    'sec_title': get_str_dict(self.project_obj.secret, 'secret'),
                                    'duty_person': self.project_obj.duty_person,
                                    'member': member,
                                    'round_num': str(int(hround.key) + 1),
                                    'round_num_chn': cname,
                                    'soft_ident': so_dut.ref,
                                    'soft_version': so_dut.version,
                                } | DocTime(payload.id).hjl_final_time(hround.key)

            result = generate_temp_doc('hjl', payload.id, round_num=cname, frag_list=payload.frag)
            if isinstance(result, dict):
                return ChenResponse(status=400, code=400,
                                    message=result.get('msg', '回归测试记录生成错误!'))
            hjl_replace_path, hjl_seitai_final_path = result
            text_frag_name_list, doc_docx = get_jinja_stdContent_element(hjl_replace_path)
            # 文本片段操作
            self.text_frag_replace_handle(text_frag_name_list, doc_docx)
            # 重新设置时序图大小(注意不变)
            for shape in doc_docx.inline_shapes:
                set_shape_size(shape)
            try:
                doc_docx.save(hjl_seitai_final_path)
            except PermissionError as e:
                return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))
        if len(cname_list) == 1:
            return get_file_respone(payload.id, '第二轮回归测试记录')
        else:
            return get_file_respone(payload.id, list(map(lambda x: f"第{x}轮回归测试说明", cname_list)))

    @route.post('/wtdDocument', url_name='create-wtdDocument')
    @transaction.atomic
    def create_wtdDocument(self, payload: SeitaiInputSchema):
        """生成最后的问题单"""
        self.project_obj = get_object_or_404(Project, id=payload.id)
        # seitai文档所需变量
        member = self.project_obj.member[0] if len(
            self.project_obj.member) > 0 else self.project_obj.duty_person
        is_jd = True if self.project_obj.report_type == '9' else False
        self.temp_context = {
                                "project_name": self.project_obj.name,
                                'project_ident': self.project_obj.ident,
                                'is_jd': is_jd,
                                'member': member,
                                'duty_person': self.project_obj.duty_person,
                                'sec_title': get_str_dict(self.project_obj.secret, 'secret'),
                            } | DocTime(payload.id).wtd_final_time()
        result = generate_temp_doc('wtd', payload.id, frag_list=payload.frag)
        if isinstance(result, dict):
            return ChenResponse(status=400, code=400,
                                message=result.get('msg', 'wtd未报出错误原因，反正在生成文档出错'))
        wtd_replace_path, wtd_seitai_final_path = result
        text_frag_name_list, doc_docx = get_jinja_stdContent_element(wtd_replace_path)
        # 文本片段操作
        self.text_frag_replace_handle(text_frag_name_list, doc_docx)
        try:
            doc_docx.save(wtd_seitai_final_path)
            return get_file_respone(payload.id, '问题单')
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

    @route.post('/bgDocument', url_name='create-bgDocument')
    @transaction.atomic
    def create_bgDocument(self, payload: SeitaiInputSchema):
        """生成最后的报告文档"""
        self.project_obj = get_object_or_404(Project, id=payload.id)
        # seitai文档所需变量
        ## 1.判断是否为JD
        member = self.project_obj.member[0] if len(
            self.project_obj.member) > 0 else self.project_obj.duty_person
        is_jd = True if self.project_obj.report_type == '9' else False
        self.temp_context = {
                                'project_name': self.project_obj.name,
                                'project_ident': self.project_obj.ident,
                                'test_purpose': "装备鉴定和列装定型" if is_jd else "软件交付和使用",
                                'is_jd': is_jd,
                                'sec_title': get_str_dict(self.project_obj.secret, 'secret'),
                                'duty_person': self.project_obj.duty_person,
                                'jd_or_third': "鉴定" if is_jd else "第三方",
                                'entrust_unit': self.project_obj.entrust_unit,
                                'member': member,
                                'joined_part': f'驻{self.project_obj.dev_unit}军事代表室、{self.project_obj.dev_unit}',
                            } | DocTime(payload.id).bg_final_time()
        result = generate_temp_doc('bg', payload.id, frag_list=payload.frag)
        if isinstance(result, dict):
            return ChenResponse(status=400, code=400,
                                message=result.get('msg', 'bg未报出错误原因，反正在生成文档出错'))
        bg_replace_path, bg_seitai_final_path = result
        text_frag_name_list, doc_docx = get_jinja_stdContent_element(bg_replace_path)
        # 文本片段操作
        self.text_frag_replace_handle(text_frag_name_list, doc_docx)
        try:
            doc_docx.save(bg_seitai_final_path)
            return get_file_respone(payload.id, '测评报告')
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

    # ~~~~模版设计模式~~~~
    # 1.传入sdtContent列表，和替换后的文档对象，进行替换操作
    def text_frag_replace_handle(self, text_frag_name_list, doc_docx: Document):
        for text_frag in text_frag_name_list:
            alias = text_frag['alias']
            if alias in self.temp_context:
                sdtContent = text_frag['sdtContent']
                stdContent_modify(self.temp_context[alias], doc_docx, sdtContent)
            else:
                print('未查找的文本片段变量：', alias)

    # ~~~~下面是生成文档辅助文本片段取变量，统一设置报错信息等，后续看重复代码封装~~~~
    # self拥有变量：self.project_obj / self.temp_context(用于替换文本片段的字典)
    # 1.获取项目第一轮round的源代码dut的用户标识/版本/第一轮测试地点
    def get_first_round_code_ident(self):
        round_obj = self.project_obj.pField.filter(key='0').first()
        if round_obj:
            self.temp_context.update({
                'location': round_obj.location,
            })
            code_dut_obj = round_obj.rdField.filter(type='SO').first()
            if code_dut_obj:
                self.temp_context.update({
                    'soft_ident': code_dut_obj.ref,
                    'soft_version': code_dut_obj.version,
                })
                return
        raise HttpError(500, "第一轮次未创建，或第一轮动态测试地点为填写，或源代码被测件未创建，请先创建")

    # 2.获取第一轮次需求规格说明dut
    def get_xq_doc_informations(self):
        round1_xq_dut = self.project_obj.pdField.filter(round__key='0', type='XQ').first()
        if round1_xq_dut:
            self.temp_context.update({'xq_version': round1_xq_dut.version})
            return
        raise HttpError(500, "第一轮次被测件：需求规格说明可能未创建，生成文档失败")

# documentType - 对应的目录名称
documentType_to_dir = {
    '测评大纲': '',
    '测试说明': 'sm',
    '测试记录': 'jl',
    '回归测试说明': 'hsm',
    '回归测试记录': 'hjl',
    '问题单': 'wtd',
    '测评报告': 'bg'
}

# 处理文档片段相关请求
@api_controller('/createfragment', tags=['生成文档-文档片段接口集合'])
class CreateFragmentController(ControllerBase):
    @route.get("/get_fragments", url_name='get-fragments')
    def get_fragements(self, id: int, documentType: str):
        """根据项目id和文档类型获取有哪些文档片段"""
        # 获取文档片段的字符串列表
        frags = self.get_fragment_name_by_document_name(id, documentType)
        # 如果没有文档片段-说明没有生成二段文档
        if not frags:
            return ChenResponse(status=500, code=500,
                                message='文档片段还未生成，请关闭后再打开/或者先下载基础文档')
        # 到这里说fragments_files数组有值，返回文件名数组
        return ChenResponse(data=[fragment for fragment in frags], message='返回文档片段成功')

    @staticmethod
    def get_fragment_name_by_document_name(id: int, document_name: str):
        # 1.找到模版的路径 - 不用异常肯定存在
        document_path = main_download_path / project_path(
            id) / 'form_template' / 'products' / f"{document_name}.docx"
        # 2.识别其中的文档片段
        frag_list = get_frag_from_document(document_path)
        # 3.这里处理报告里第十轮次前端展示问题
        ## 3.1先判断是否为报告
        if document_name == '测评报告':
            ## 3.2然后判断有几个轮次
            project_obj = get_object_or_404(Project, id=id)
            round_qs = project_obj.pField.all()
            white_list_frag = []
            ## 3.3将希望有的片段名称加入白名单
            for round_obj in round_qs:
                chn_num = digit_to_chinese(int(round_obj.key) + 1)
                exclude_str = f"测试内容和结果_第{chn_num}轮次"  # 组成识别字符串
                white_list_frag.append(exclude_str)
            ## 3.4过滤包含“测试内容和结果的轮次在白名单的通过”
            # 去掉所有“测试内容和结果_”的片段
            filter_frags = list(filter(lambda x: '测试内容和结果' not in x['frag_name'], frag_list))
            # 再找到白名单的“测试内容和结果_”的片段
            content_and_result_frags = list(
                filter(lambda x: '测试内容和结果' in x['frag_name'] and x['frag_name'] in white_list_frag,
                       frag_list))
            # 再组合起来返回
            filter_frags.extend(content_and_result_frags)
            return filter_frags
        return frag_list

    @route.get("/get_round_exit", url_name='get-round-exit')
    def get_round_exit(self, id: int):
        """该函数主要识别有几轮回归测试说明、几轮回归测试记录"""
        project_obj: Project = get_object_or_404(Project, id=id)
        # 取非第一轮次的轮次的个数
        round_count = project_obj.pField.exclude(key='0').count()
        return {'count': round_count}

# 自定义修改Django的文件系统-启动覆盖模式
class OverwriteStorage(FileSystemStorage):
    def __init__(self, *args, **kwargs):
        kwargs['allow_overwrite'] = True  # 启用覆盖模式
        super().__init__(*args, **kwargs)

def digit_to_chinese(num):
    num_dict = {'0': '零', '1': '一', '2': '二', '3': '三', '4': '四',
                '5': '五', '6': '六', '7': '七', '8': '八', '9': '九', '10': '十'}
    return ''.join(num_dict[d] for d in str(num))

# 处理用户上传有文档片段的产品文档文件：注意回归测试说明、回归测试记录需要单独处理
@api_controller('/documentUpload', tags=['生成文档-上传模版文档接口'])
class UploadDocumentController(ControllerBase):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # 储存上传文件
        self.upload_file: UploadedFile | None = None

    @route.post("/file", url_name='upload-file')
    def upload_file(self, id: int, documentType: str, file: File[UploadedFile], round_num: int = None):
        self.upload_file = file
        # 1.获取储存路径
        target_dir = main_download_path / project_path(id) / 'form_template' / 'products'
        # 2.初始化文件系统
        fs = OverwriteStorage(location=target_dir)

        # 新：如果是大纲片段，则大纲所有片段以文档片段方式储存在/reuse文件夹下面
        if documentType == '测评大纲':
            self.get_dg_to_reuse_dir(target_dir.parent.parent / 'reuse')

        if round_num is None:
            # 处理非“回归测试说明”/“回归测试记录”文档的上传
            # warning：不校验文档内是否有文档片段，由用户保证上传内容
            # 3.覆盖储存，注意会返回文件的name属性
            fs.save(f"{documentType}.docx", self.upload_file)
        else:
            # 处理“回归测试说明”/“回归测试记录”文档的上传
            fs.save(f"第{digit_to_chinese(round_num)}轮{documentType}.docx", self.upload_file)
        return ChenResponse(status=200, code=200, message=f'上传{documentType}成功！')

    # 主功能函数：将所有大纲的片段储存在reuse下面，以便其他文件使用
    def get_dg_to_reuse_dir(self, reuse_dir_path: Path):
        """将大纲的文档片段储存在/reuse文件夹下面"""
        doc = Document(self.upload_file)
        frag_list = self.get_document_frag_list(doc)
        for frag_item in frag_list:
            # 目的是格式明确按照“测评大纲.docx”进行，后续文档一样必须按照这样
            new_doc = Document((reuse_dir_path / 'basic_doc.docx').as_posix())
            if frag_item['content'] is not None:
                # XML元素可以直接append
                new_doc.element.body.clear_content()
                for frag_child in frag_item['content'].iterchildren():
                    new_doc.element.body.append(frag_child)
            filename = f"{frag_item['alias']}.docx"
            new_doc.save((reuse_dir_path / filename).as_posix())

    # 辅助函数：将上传文件的文档片段以列表形式返回
    def get_document_frag_list(self, doc: Document):
        body = doc.element.body
        sdt_element_list = body.xpath('./w:sdt')  # 只查询文档片段，非文本片段
        frag_list = []
        for sdt_element in sdt_element_list:
            alias_name = None
            sdtContent = None
            for sdt_child in sdt_element.iterchildren():
                if sdt_child.tag.endswith('sdtPr'):
                    for sdtPr_child in sdt_child.getchildren():
                        if sdtPr_child.tag.endswith('alias'):
                            if len(sdtPr_child.attrib.values()) > 0:
                                alias_name = sdtPr_child.attrib.values()[0]
                if sdt_child.tag.endswith("sdtContent"):
                    sdtContent = sdt_child
            frag_list.append({'alias': alias_name, 'content': sdtContent})
        return list(filter(lambda x: x['alias'] is not None, frag_list))
