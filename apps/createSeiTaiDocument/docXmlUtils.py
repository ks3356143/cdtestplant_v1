"""该文件是：替换文档片段然后生成辅助生成最终文档"""
from io import BytesIO
from typing import List, Dict
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.oxml.shape import CT_Picture
from docx.parts.image import ImagePart
from docx.text.run import Run
from docx.shared import Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml.etree import _Element

# 路径工具
from utils.path_utils import project_path

### 模块变量：定义常用图片所在区域的宽高
Demand_table_xqms = Mm(134)  # 1.测评大纲-测试项里面-需求描述单元格
Timing_diagram_width = Mm(242)  # 2.测试记录-时序图
Test_result_width = Mm(78)  # 3.测试记录-测试结果
Horizatal_width = Mm(130)  # 4.所有文档-页面图片的横向距离（图片宽度预设置）

def getParentRunNode(node):
    """传入oxml节点对象，获取其祖先节点的CT_R"""
    if isinstance(node, CT_R):
        return node
    return getParentRunNode(node.getparent())

def generate_temp_doc(doc_type: str, project_id: int, round_num=None, frag_list=None):
    """ 该函数参数：
    :param frag_list: 储存用户不覆盖的片段列表
    :param round_num: 只有回归说明和回归记录有
    :param project_id: 项目id
    :param doc_type:大纲 sm:说明 jl:记录 bg:报告 hsm:回归测试说明 hjl:回归测试记录,默认路径为dg -> 所以如果传错就生成生成大纲了
    :return (to_tpl_file路径, seitai_final_file路径)
    """
    if frag_list is None:
        frag_list = []
    # ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
    project_path_str = project_path(project_id)
    # 根据传入需要处理的文档类型，自动获路径
    prefix = Path.cwd() / 'media' / project_path_str
    template_file: Path = prefix / 'form_template' / 'products' / '测评大纲.docx'
    to_tpl_file: Path = prefix / 'temp' / '测评大纲.docx'
    seitai_final_file: Path = prefix / 'final_seitai' / '测评大纲.docx'
    if doc_type == 'sm':
        template_file = prefix / 'form_template' / 'products' / '测试说明.docx'
        to_tpl_file = prefix / 'temp' / '测试说明.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / '测试说明.docx'
    elif doc_type == 'jl':
        template_file = prefix / 'form_template' / 'products' / '测试记录.docx'
        to_tpl_file = prefix / 'temp' / '测试记录.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / '测试记录.docx'
    elif doc_type == 'bg':
        template_file = prefix / 'form_template' / 'products' / '测评报告.docx'
        to_tpl_file = prefix / 'temp' / '测评报告.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / '测评报告.docx'
    elif doc_type == 'hsm':
        # 如果products里面存在“用户上传的第n轮回归测试说明.docx，则使用它作为模版”
        template_file = prefix / 'form_template' / 'products' / f'第{round_num}轮回归测试说明.docx'
        if not template_file.exists():
            template_file = prefix / 'form_template' / 'products' / '回归测试说明.docx'
        to_tpl_file = prefix / 'temp' / f'第{round_num}轮回归测试说明.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / f'第{round_num}轮回归测试说明.docx'
    elif doc_type == 'hjl':
        # 如果products里面存在“用户上传的第n轮回归测试记录.docx，则使用它作为模版”
        template_file = prefix / 'form_template' / 'products' / f'第{round_num}轮回归测试记录.docx'
        if not template_file.exists():
            template_file = prefix / 'form_template' / 'products' / '回归测试记录.docx'
        to_tpl_file = prefix / 'temp' / f'第{round_num}轮回归测试记录.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / f'第{round_num}轮回归测试记录.docx'
    elif doc_type == 'wtd':
        template_file = prefix / 'form_template' / 'products' / '问题单.docx'
        to_tpl_file = prefix / 'temp' / '问题单.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / '问题单.docx'
    # 定义找寻被复制文件根路径 - 后续会根据type找子路径
    output_files_path = prefix / 'output_dir'
    # 这里可能修改，储存大纲里面的文档片段
    dg_copied_files = []
    # 储存sm/jl/hsm/hjl/bg/wtd的文档片段
    exclusive_copied_files = []
    # 新：储存reuse的文档片段
    reuse_files = []
    # 将被拷贝文件分别放入不同两个数组
    for file in output_files_path.iterdir():
        if file.is_file():
            if file.suffix == '.docx':
                dg_copied_files.append(file)
        elif file.is_dir():
            # 如果文件夹名称为sm/jl/hsm/hjl/bg/wtd则进入该判断
            # 所以要求文件系统文件夹名称必须是sm/jl/hsm/hjl/bg/wtd不然无法生成
            if file.stem == doc_type:
                for f in file.iterdir():
                    if f.suffix == '.docx':
                        exclusive_copied_files.append(f)
    for file in (prefix / 'reuse').iterdir():
        if file.is_file():
            if file.suffix == '.docx':
                reuse_files.append(file)
    # 找到基础模版的所有std域
    doc = Document(template_file.as_posix())
    body = doc.element.body
    sdt_element_list = body.xpath('./w:sdt')
    # 找到sdt域的名称 -> 为了对应output_dir文件 / 储存所有output_dir图片
    area_name_list = []
    image_part_list = []  # 修改为字典两个字段{ 'name':'测评对象', 'img':ImagePart }
    # 筛选片段【二】：用户前端要求不要覆盖的文档片段
    frag_is_cover_dict = {item.name: item.isCover for item in frag_list}
    # 遍历所有控件 -> 放入area_name_list【这里准备提取公共代码】
    for sdt_ele in sdt_element_list:
        isLock = False
        for elem in sdt_ele.iterchildren():
            # 【一】用户设置lock - 下面2个if将需要被替换的(控件名称)存入area_name_list
            if elem.tag.endswith('sdtPr'):
                for el in elem.getchildren():
                    if el.tag.endswith('lock'):
                        isLock = True
            if elem.tag.endswith('sdtPr'):
                for el in elem.getchildren():
                    if el.tag.endswith('alias'):
                        # 筛序【一】：取出用户设置lock的文档片段
                        if len(el.attrib.values()) > 0 and (isLock == False):
                            area_name = el.attrib.values()[0]
                            # 筛选【二】：前端用户选择要覆盖的片段
                            if frag_is_cover_dict.get(area_name):
                                area_name_list.append(area_name)
            # 下面开始替换area_name_list的“域”（这时已经被筛选-因为sdtPr和sdtContent是成对出现）
            if elem.tag.endswith('sdtContent'):
                if len(area_name_list) > 0:
                    # 从第一个片段名称开始取，取到模版的“域”名称
                    area_pop_name = area_name_list.pop(0)
                    # 这里先去找media/output_dir/xx下文件，然后找media/output下文件
                    copied_file_path = ""
                    # 下面if...else是找output_dir下面文件与“域”名称匹配，匹配到存入copied_file_path
                    if doc_type == 'dg':
                        for file in dg_copied_files:
                            if file.stem == area_pop_name:
                                copied_file_path = file
                    else:
                        # 如果不是大纲
                        if round_num is None:
                            # 如果非回归说明、记录
                            for file in exclusive_copied_files:
                                if file.stem == area_pop_name:
                                    copied_file_path = file
                            # 这里判断是否copied_file_path没取到文件，然后遍历reuse下文件
                            if not copied_file_path:
                                for file in reuse_files:
                                    if file.stem == area_pop_name:
                                        copied_file_path = file
                            # 如果上面被复制文件还没找到，然后遍历output_dir下文件
                            if not copied_file_path:
                                for file in dg_copied_files:
                                    if file.stem == area_pop_name:
                                        copied_file_path = file
                        else:
                            # 因为回归的轮次，前面会加 -> 第{round_num}轮
                            for file in exclusive_copied_files:  # 这里多了第{round_num}轮
                                if file.stem == f"第{round_num}轮{area_pop_name}":
                                    copied_file_path = file
                            if not copied_file_path:
                                for file in reuse_files:
                                    if file.stem == area_pop_name:
                                        copied_file_path = file
                            if not copied_file_path:
                                for file in dg_copied_files:
                                    if file.stem == area_pop_name:
                                        copied_file_path = file
                    # 找到文档片段.docx，将其数据复制到对应area_name的“域”
                    if copied_file_path:
                        doc_copied = Document(copied_file_path)
                        copied_element_list = []
                        element_list = doc_copied.element.body.inner_content_elements
                        for elet in element_list:
                            if isinstance(elet, CT_P):
                                copied_element_list.append(Paragraph(elet, doc_copied))
                            if isinstance(elet, CT_Tbl):
                                copied_element_list.append(Table(elet, doc_copied))
                        elem.clear()
                        for para_copied in copied_element_list:
                            elem.append(para_copied._element)

                        # 下面代码就是将图片全部提取到image_part_list，以便后续插入，注意这时候已经是筛选后的
                        doc_copied = Document(copied_file_path)  # 需要重新获取否则namespace错误
                        copied_body = doc_copied.element.body
                        img_node_list = copied_body.xpath('.//pic:pic')
                        if not img_node_list:
                            pass
                        else:
                            for img_node in img_node_list:
                                img: CT_Picture = img_node
                                # 根据节点找到图片的关联id
                                embed = img.xpath('.//a:blip/@r:embed')[0]
                                # 这里得到ImagePart -> 马上要给新文档添加
                                related_part: ImagePart = doc_copied.part.related_parts[embed]
                                # doc_copied.part.related_parts是一个字典
                                image_part_list.append({'name': area_pop_name, 'img': related_part})

    # 现在是替换后，找到替换后文档所有pic:pic，并对“域”名称进行识别
    graph_node_list = body.xpath('.//pic:pic')
    graph_node_list_transform = []
    for picNode in graph_node_list:
        # 遍历替换后模版的所有pic，去找祖先
        sdt_node = picNode.xpath('ancestor::w:sdt[1]')[0]
        for sdt_node_child in sdt_node.iterchildren():
            # 找到sdt下一级的stdPr
            if sdt_node_child.tag.endswith('sdtPr'):
                for sdtPr_node_child in sdt_node_child.getchildren():
                    if sdtPr_node_child.tag.endswith('alias'):
                        yu_name = sdtPr_node_child.attrib.values()[0]
                        graph_node_list_transform.append({'yu_name': yu_name, 'yu_node': picNode})
    for graph_node in graph_node_list_transform:
        image_run_node = getParentRunNode(graph_node['yu_node'])
        image_run_node.clear()
        # 循环去image_part_list找name和yu_name相等的图片
        for img_part in image_part_list:
            # 1.如果找到相等
            if img_part['name'] == graph_node['yu_name']:
                # 2.找到即可添加图片到“域”
                image_run_node.clear()
                # 辅助：去找其父节点是否为段落，是段落则存起来，后面好居中
                image_run_parent_paragraph = image_run_node.getparent()
                father_paragraph = None
                if isinstance(image_run_parent_paragraph, CT_P):
                    father_paragraph = Paragraph(image_run_parent_paragraph, doc)
                copied_bytes_io = BytesIO(img_part['img'].image.blob)
                r_element = Run(image_run_node, doc)
                inline_shape = r_element.add_picture(copied_bytes_io)
                ## 2.1.统一：这里设置文档片段里面的图片大小和位置
                source_width = inline_shape.width
                source_height = inline_shape.height
                if source_width >= source_height:
                    inline_shape.width = Mm(120)
                    inline_shape.height = int(inline_shape.height * (inline_shape.width / source_width))
                else:
                    inline_shape.height = Mm(60)
                    inline_shape.width = int(inline_shape.width * (inline_shape.height / source_height))
                ## 2.2.设置图片所在段落居中对齐
                if father_paragraph:
                    father_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                r_element.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # 3.因为按顺序的，所以移除image_part_list中已经替换的图片
                image_part_list.remove(img_part)
                break
    try:
        # 这里直接生成产品文档
        doc.save(str(to_tpl_file))
        return to_tpl_file, seitai_final_file
    except PermissionError as e:
        return {'code': 'error', 'msg': '生成的temp文件已打开，请关闭后重试...'}

def get_frag_from_document(doc_path: Path) -> List[Dict]:
    """传入products的文件路径，识别出所有文档片段名称，数组返回：要求docx里面文档名称不能更变"""
    doc = Document(doc_path.as_posix())
    sdt_element_list = doc.element.body.xpath('./w:sdt')
    # 整个for循环识别文档片段名称
    area_name_list = []
    for sdt_ele in sdt_element_list:
        isLock = False
        alias_value = None
        for elem in sdt_ele.iterchildren():
            if elem.tag.endswith('sdtPr'):
                for el in elem.getchildren():
                    if el.tag.endswith('alias'):
                        alias_value = el.attrib.values()
                    # 查找是否被用户在模版上标记了Lock
                    if el.tag.endswith('lock'):
                        isLock = True
        if alias_value and len(alias_value):
            area_name_list.append({'frag_name': alias_value[0], 'isLock': isLock})
    return area_name_list

# 辅助函数-传入temp文件路径（已替换文档片段的temp文档），输出stdContent
def get_jinja_stdContent_element(temp_docx_path: Path):
    doc_docx = Document(temp_docx_path.as_posix())
    body = doc_docx.element.body
    # 储存文本片段
    text_frag_name_list = []
    sdt_element_list = body.xpath('//w:sdt')

    # 注意python-docx的页头的文本片段不在body里面，而在section.header里面
    # 所以定义辅助函数，统一处理
    def deel_sdt_content(*args):
        """传入sdt_element列表，将其sdtContent加入外部的文本片段列表"""
        for sdt_ele in args:
            # 找出每个sdt下面的3个标签
            tag_value = None
            alias_value = None
            sdtContent_ele = None
            for sdt_ele_child in sdt_ele.iterchildren():
                if sdt_ele_child.tag.endswith('sdtPr'):
                    for sdtPr_ele_child in sdt_ele_child.getchildren():
                        if sdtPr_ele_child.tag.endswith('tag'):
                            if len(sdtPr_ele_child.attrib.values()) > 0:
                                tag_value = sdtPr_ele_child.attrib.values()[0]
                        if sdtPr_ele_child.tag.endswith('alias'):
                            if len(sdtPr_ele_child.attrib.values()) > 0:
                                alias_value = sdtPr_ele_child.attrib.values()[0]
                if sdt_ele_child.tag.endswith('sdtContent'):
                    sdtContent_ele = sdt_ele_child
            # 找出所有tag_value为jinja的文本片段
            if tag_value == 'jinja' and alias_value is not None and sdtContent_ele is not None:
                text_frag_name_list.append({'alias': alias_value, 'sdtContent': sdtContent_ele})

    deel_sdt_content(*sdt_element_list)
    for section in doc_docx.sections:
        header = section.header
        header_sdt_list = header.part.element.xpath('//w:sdt')
        deel_sdt_content(*header_sdt_list)

    return text_frag_name_list, doc_docx

# 封装一个根据alias名称修改stdContent的函数 -> 在接口处理函数中取数据放入函数修改文档
def stdContent_modify(modify_str: str | bool, doc_docx: Document, sdtContent: _Element):
    # 正常处理
    for ele in sdtContent:
        if isinstance(ele, CT_R):
            run_ele = Run(ele, doc_docx)
            if isinstance(modify_str, bool):
                # 如果是True，则不修改原来
                if modify_str:
                    break
                else:
                    modify_str = ""
            # 有时候会int类型，转换一下防止报错
            if isinstance(modify_str, int):
                modify_str = str(modify_str)
            run_ele.text = modify_str
            sdtContent.clear()
            sdtContent.append(run_ele._element)
            break

        if isinstance(ele, CT_P):
            para_ele = Paragraph(ele, doc_docx)
            if isinstance(modify_str, bool):
                if modify_str:
                    break
                else:
                    modify_str = ""
            para_ele.clear()
            para_ele.text = modify_str
            sdtContent.clear()
            sdtContent.append(para_ele._element)
            break
