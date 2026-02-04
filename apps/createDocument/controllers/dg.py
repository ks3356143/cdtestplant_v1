import base64
import io
from typing import Any
from datetime import datetime
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from ninja.errors import HttpError
from ninja_extra import ControllerBase, api_controller, route
from django.db import transaction
from django.db.models import Q
from docxtpl import DocxTemplate, InlineImage
from pathlib import Path
from utils.chen_response import ChenResponse
# 导入数据库ORM
from apps.project.models import Project, Contact, Abbreviation, ProjectSoftSummary, StuctSortData
from apps.dict.models import Dict
# 导入工具函数
from utils.util import get_str_dict, get_list_dict, get_testType, get_ident, get_str_abbr
from utils.chapter_tools.csx_chapter import create_csx_chapter_dict
from django.shortcuts import get_object_or_404
from django.forms.models import model_to_dict
from apps.createDocument.extensions.util import create_dg_docx
from apps.createDocument.extensions.parse_rich_text import RichParser
from apps.createDocument.extensions.documentTime import DocTime
from utils.path_utils import project_path
# 记录生成日志
from apps.createSeiTaiDocument.extensions.logger import GenerateLogger
# 导入mixins-处理文档片段
from apps.createDocument.extensions.mixins import FragementToolsMixin
# 导入工具
from apps.createDocument.extensions.tools import demand_sort_by_designKey, set_table_border

# @api_controller("/generate", tags=['生成大纲文档'], auth=JWTAuth(), permissions=[IsAuthenticated])
@api_controller("/generate", tags=['生成大纲文档'])
class GenerateControllerDG(ControllerBase, FragementToolsMixin):
    logger = GenerateLogger('测评大纲')

    @route.get("/create/testdemand", url_name="create-testdemand")
    @transaction.atomic
    def create_testdemand(self, id: int):  # type:ignore
        """目前生成第一轮测试项"""
        tplTestDemandGenerate_path = Path.cwd() / "media" / project_path(
            id) / "form_template" / "dg" / "测试项及方法.docx"
        doc = DocxTemplate(tplTestDemandGenerate_path)
        # 获取指定的项目对象
        project_qs = get_object_or_404(Project, id=id)
        # 先查询dict字典，查出总共有多少个testType
        test_type_len = Dict.objects.get(code='testType').dictItem.count()
        type_number_list = [i for i in range(1, test_type_len + 1)]
        list_list = [[] for _ in range(1, test_type_len + 1)]

        # 查出第一轮所有testdemand
        project_round_one = project_qs.pField.filter(key=0).first()
        testDemand_qs = project_round_one.rtField.all().select_related('design')
        # 按照自己key排序，这样可以按照design的key排序
        sorted_demand_qs = sorted(testDemand_qs, key=demand_sort_by_designKey)

        # 遍历第一轮测试项：默认是ID排序
        for single_qs in sorted_demand_qs:
            type_index = type_number_list.index(int(single_qs.testType))
            # 先查询其testDemandContent信息
            content_list = []
            for (index, content) in enumerate(single_qs.testQField.all()):
                content_dict = {
                    "index": index + 1,
                    "rindex": str(index + 1).rjust(2, '0'),
                    "subName": content.subName,
                    "subDescription": content.subDescription,
                    # 修改遍历content下面的step，content变量是TestDemandContent表
                    "subStep": [
                        {'index': index + 1, 'operation': step_obj.operation, 'expect': step_obj.expect}
                        for (index, step_obj) in enumerate(content.testStepField.all())
                    ],
                }
                content_list.append(content_dict)
            # 查询测试项中testMethod
            testmethod_str = ''
            for dict_item_qs in Dict.objects.get(code="testMethod").dictItem.all():
                for tm_item in single_qs.testMethod:
                    if tm_item == dict_item_qs.key:
                        testmethod_str += dict_item_qs.title + " "
            # 富文本解析
            # ***Inspect-start：检查设计需求的描述是否为空***
            if single_qs.design.description == '':
                design_info = single_qs.design.ident + '-' + single_qs.design.name
                self.logger.write_warning_log('测试项', f'设计需求中的描述为空，请检查 -> {design_info}')
            # ***Inspect-end***
            html_parser = RichParser(single_qs.design.description)
            desc_list = html_parser.get_final_list(doc)
            # 查询关联design以及普通design
            doc_list = [{'dut_name': single_qs.dut.name, 'design_chapter': single_qs.design.chapter,
                         'design_name': single_qs.design.name}]
            for relate_design in single_qs.otherDesign.all():
                ddict = {'dut_name': relate_design.dut.name, 'design_chapter': relate_design.chapter,
                         'design_name': relate_design.name}
                doc_list.append(ddict)

            # 组装单个测试项
            ## 打印本项目是FPGA还是CPU
            testdemand_dict = {
                "name": single_qs.name,
                "key": single_qs.key,
                "ident": get_ident(single_qs),
                "priority": get_str_dict(single_qs.priority, "priority"),
                "doc_list": doc_list,
                "design_description": desc_list,
                "test_demand_content": content_list,
                "testMethod": testmethod_str.strip(),
                "adequacy": single_qs.adequacy.replace("\n", "\a"),
                "testDesciption": single_qs.testDesciption.replace("\n", "\a"),  # 测试项描述
                "testType": get_testType(single_qs.testType, 'testType'),
            }
            list_list[type_index].append(testdemand_dict)

        output_list = []

        for (index, li) in enumerate(list_list):
            qs = Dict.objects.get(code="testType").dictItem.get(key=str(index + 1))
            context_str = qs.title
            sort = qs.sort
            table = {
                "type": context_str,
                "item": li,
                "sort": sort
            }
            output_list.append(table)

        # 排序1：测试类型排序
        output_list = sorted(output_list, key=(lambda x: x["sort"]))

        # 定义渲染context字典
        context = {
            "project_name": project_qs.name,
            "is_JD": True if project_qs.report_type == '9' else False,
            "data": output_list,
            "isFPGA": '1' in project_qs.plant_type
        }

        doc.render(context, autoescape=True)
        try:
            doc.save(Path.cwd() / "media" / project_path(id) / "output_dir" / "测试项及方法.docx")
            return ChenResponse(status=200, code=200, message="文档生成成功！")
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

    @route.get("/create/yiju", url_name='create-yiju')
    @transaction.atomic
    def create_yiju(self, id: int):
        # 先找出所属项目
        project_qs = get_object_or_404(Project, id=id)
        # 找出该项目的真实依据文件qs
        yiju_list = get_list_dict('standard', project_qs.standard)
        context = {
            'std_documents': yiju_list
        }
        return create_dg_docx('标准依据文件.docx', context, id)

    @route.get("/create/techyiju", url_name='create-techyiju')
    @transaction.atomic
    def create_techyiju(self, id: int):
        # 找出所属项目
        project_qs = get_object_or_404(Project, id=id)
        # 根据项目找出被测件-只找第一轮次
        duties_qs = project_qs.pdField.filter(
            Q(type='XQ') | Q(type='SJ') | Q(type='XY') | Q(type='YZ')).filter(
            round__key='0')
        # 先定义个字典
        std_documents = []
        for duty in duties_qs:
            one_duty = {'doc_name': duty.name, 'ident_version': duty.ref + '-' + duty.version,
                        'publish_date': duty.release_date, 'source': duty.release_union}
            std_documents.append(one_duty)

        # 生成二级文档
        context = {
            'std_documents': std_documents
        }
        return create_dg_docx('技术依据文件.docx', context, id)

    @route.get("/create/contact", url_name='create-contact')
    @transaction.atomic
    def create_contact(self, id: int):
        # 先找出所属项目
        project_qs = get_object_or_404(Project, id=id)
        contact_dict = model_to_dict(project_qs,
                                     fields=['entrust_unit', 'entrust_contact', 'entrust_contact_phone',
                                             'dev_unit',
                                             'dev_contact', 'dev_contact_phone', 'test_unit', 'test_contact',
                                             'test_contact_phone'])
        # 根据entrust_unit、dev_unit、test_unit查找Contact中地址信息
        entrust_addr = Contact.objects.get(name=contact_dict['entrust_unit']).addr
        dev_addr = Contact.objects.get(name=contact_dict['dev_unit']).addr
        test_addr = Contact.objects.get(name=contact_dict['test_unit']).addr
        contact_dict['entrust_addr'] = entrust_addr
        contact_dict['dev_addr'] = dev_addr
        contact_dict['test_addr'] = test_addr
        context = {
            'datas': contact_dict
        }
        return create_dg_docx('联系人和方式.docx', context, id)

    # 生成测评时间和地点
    @route.get('/create/timeaddress', url_name='create-timeaddress')
    @transaction.atomic
    def create_timeaddress(self, id: int):
        doc_timer = DocTime(id)
        context = doc_timer.dg_address_time()
        context = self.change_time_to_another(context, ['beginTime_strf', 'dgCompileStart', 'dgCompileEnd',
                                                        'designStart', 'designEnd'])
        return create_dg_docx('测评时间和地点.docx', context, id)

    # 2025/12/11：将20250417格式改为2025年04月17日 - 封装函数，传入字典和键值，修改对应键值信息
    def change_time_to_another(self, context: dict, key_list: list[str]):
        for key in key_list:
            time_val = context.get(key, None)
            if time_val:
                context[key] = datetime.strptime(time_val, "%Y%m%d").strftime("%Y年%m月%d日")
        return context

    # 生成【主要功能和性能指标】文档片段
    @route.get('/create/indicators', url_name='create-indicators')
    @transaction.atomic
    def create_indicators(self, id: int):
        # 获取文档片段模版路径
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '主要功能和性能指标.docx'
        doc = DocxTemplate(input_path)
        # 获取项目对象
        project_obj: Project = get_object_or_404(Project, id=id)
        # 定义JINJA上下文
        # 获取第一轮次所有功能、性能的设计需求[目前只支持XQ和YZ]，因为要获取dut信息进行连表查询
        q_ex = Q(dut__type='XQ') | Q(dut__type='YZ')
        design_qs = project_obj.psField.filter(q_ex, round__key='0').select_related('dut')
        # 1.功能性能覆盖表
        # 定义功能/性能两个列表
        func_design_list = []
        performance_design_list = []
        # 遍历设计需求，然后放入两个列表
        for design_obj in design_qs:
            # 功能指标描述
            description = RichParser(design_obj.description).get_final_p_list()
            # 覆盖情况-【查询测试项-测试子项名称】
            demand_qs = design_obj.dtField.all()
            str_list = []
            for demand in demand_qs:
                # 再查询测试子项
                for subDemand in demand.testQField.all():
                    str_list.append(subDemand.subName)
            coverage_str = "、".join(str_list)
            design_context_obj = {
                'chapter_info': f"《{design_obj.dut.name}》{design_obj.chapter}-{design_obj.name}",
                'indicator': "\a".join(description),
                'coverage': f"对{design_obj.name}进行全覆盖测试，包含{coverage_str}，验证所描述内容是否满足需求等文档的要求"
            }
            demandType_str = get_str_dict(design_obj.demandType, 'demandType')
            # 判断是否包含“功能”/“性能”字样
            if '功能' in demandType_str:
                func_design_list.append(design_context_obj)
            elif '性能' in demandType_str:
                performance_design_list.append(design_context_obj)
        # 2.摸底指标清单
        # 先查第一轮次所有测试项
        is_has_modi = False
        md_demand_list = []
        round1_demand_qs = project_obj.ptField.filter(round__key='0')
        for one_demand in round1_demand_qs:
            testType_str = get_str_dict(one_demand.testType, 'testType')
            if '摸底' in testType_str:
                is_has_modi = True
                md_demand_list.append({
                    'xq_source': "隐含需求",
                    'desc': one_demand.testDesciption,
                    'demand_name': one_demand.name,
                    'demand_ident': "".join(["XQ_MD_", one_demand.ident])
                })
        # 上下文添加
        context = {
            'project_name': project_obj.name,
            'func_design_list': func_design_list,
            'performance_design_list': performance_design_list,
            'md_demand_list': md_demand_list,
            'is_has_modi': is_has_modi
        }
        doc.render(context, autoescape=True)
        try:
            doc.save(Path.cwd() / "media" / project_path(id) / "output_dir" / '主要功能和性能指标.docx')
            return ChenResponse(status=200, code=200, message="文档生成成功！")
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

    # 生成测评对象 - 包括大纲、说明、回归说明和报告
    @route.get('/create/softComposition', url_name='create-softComposition')
    @transaction.atomic
    def create_softComposition(self, id: int):
        # 首先判断是否包含 - 项目信息-软件概述
        project_obj = get_object_or_404(Project, id=id)
        input_path_2 = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '测评对象_2.docx'
        doc = DocxTemplate(input_path_2)
        soft_summary_qs = ProjectSoftSummary.objects.filter(project=project_obj)
        if soft_summary_qs.exists():
            data_qs = soft_summary_qs.first().data_schemas
            if data_qs.exists():
                # 如果存在则渲染此处
                data_list = []
                for data_obj in data_qs.all():
                    item_context: dict[str, Any] = {"fontnote": data_obj.fontnote, 'type': data_obj.type}
                    # 根据数据类型处理content字段
                    if data_obj.type == 'text':
                        item_context['content'] = data_obj.content
                    elif data_obj.type == 'table':
                        # 使用subdoc
                        subdoc = doc.new_subdoc()
                        rows = len(data_obj.content)
                        cols = len(data_obj.content[0])
                        table = subdoc.add_table(rows=rows, cols=cols, style='Table Grid')
                        # 设置边框
                        set_table_border(table)
                        # 单元格处理
                        for row in range(rows):
                            for col in range(cols):
                                cell = table.cell(row, col)
                                cell.text = data_obj.content[row][col]
                                # 第一行设置居中
                                if row == 0:
                                    # 黑体设置
                                    cell.text = ""
                                    pa = cell.paragraphs[0]
                                    run = pa.add_run(str(data_obj.content[row][col]))
                                    run.font.name = '黑体'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                                    run.font.bold = False
                                    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 表格居中
                        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        item_context['content'] = subdoc
                    elif data_obj.type == 'image':
                        base64_bytes = base64.b64decode(data_obj.content.replace("data:image/png;base64,", ""))
                        item_context['content'] = InlineImage(doc, io.BytesIO(base64_bytes), width=Mm(120))
                    data_list.append(item_context)
                context = {
                    "datas": data_list,
                }
                doc.render(context)
                try:
                    doc.save(Path.cwd() / "media" / project_path(id) / "output_dir" / '测评对象.docx')
                    return ChenResponse(status=200, code=200, message="文档生成成功！")
                except PermissionError as e:
                    return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

        # 原来文档片段或者初始内容
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '测评对象.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '测评对象')
        context = {
            "replace": replace,  # 指定是否由数据库文档片段进行生成
            "user_content": frag and rich_text_list
        }
        doc.render(context, autoescape=True)
        try:
            doc.save(Path.cwd() / "media" / project_path(id) / "output_dir" / '测评对象.docx')
            return ChenResponse(status=200, code=200, message="文档生成成功！")
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

    # 生成被测软件接口章节
    @route.get('/create/interface', url_name='create-interface')
    def create_interface(self, id: int):
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '被测软件接口.docx'
        doc = DocxTemplate(input_path)
        project_qs = get_object_or_404(Project, id=id)
        project_name = project_qs.name
        interfaceNameList = []
        # 查询接口列表
        iters = project_qs.psField.filter(demandType=3)
        iters_length = len(iters)
        index = 0
        for inter in iters:
            interfaceNameList.append(inter.name)
            index += 1
            if index < iters_length:
                interfaceNameList.append('、')
        # 对每个接口进行字典处理
        interface_list = []
        for interface in iters:
            interface_dict = {
                'name': interface.name,
                'ident': interface.ident,
                'source': interface.source,
                'to': interface.to,
                'type': interface.type,
                'protocal': interface.protocal,
            }
            interface_list.append(interface_dict)
        # 项目接口图处理 - 2026/2/4
        image_obj = StuctSortData.objects.filter(project=project_qs)
        ## 判断是否存在
        image_render = None
        fontnote = None
        if image_obj.exists():
            base64_bytes = base64.b64decode(image_obj.first().content.replace("data:image/png;base64,", ""))
            image_render = InlineImage(doc, io.BytesIO(base64_bytes), width=Mm(120))
            fontnote = image_obj.first().fontnote
        context = {
            'project_name': project_name,
            'iters': interfaceNameList,
            'iter_list': interface_list,
            'image_render': image_render if image_render else "",
            'fontnote': fontnote if fontnote else "".join([project_name, '接口示意图'])
        }
        doc.render(context, autoescape=True)
        try:
            doc.save(Path.cwd() / "media" / project_path(id) / "output_dir" / '被测软件接口.docx')
            return ChenResponse(status=200, code=200, message="文档生成成功！")
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

    # 生成顶层技术文件
    @route.get('/create/top_file', url_name='create-performance')
    def create_top_file(self, id: int):
        project_obj: Project = get_object_or_404(Project, id=id)
        is_JD = True if project_obj.report_type == '9' else False
        dut_qs = project_obj.pdField.filter(type='YZ')
        dut_list = [{
            'index': index + 2 if is_JD else index + 1,
            'name': dut_obj.name,
            'ident_and_version': '-'.join([dut_obj.ref, dut_obj.version]),
            'publish_date': dut_obj.release_date,
            'source': dut_obj.release_union,
        } for index, dut_obj in enumerate(dut_qs)]
        context = {
            'project_name': project_obj.name,
            'is_JD': is_JD,
            'dut_list': dut_list,
        }
        return create_dg_docx('顶层技术文件.docx', context, id)

    # 静态测试环境说明
    @route.get('/create/static_env', url_name='create-static_env')
    def create_static_env(self, id: int):
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '静态测试环境说明.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '静态测试环境说明')
        context = {
            "replace": replace,  # 指定是否由数据库文档片段进行生成
            "user_content": frag and rich_text_list
        }
        doc.render(context, autoescape=True)
        try:
            doc.save(Path.cwd() / "media" / project_path(id) / "output_dir" / '静态测试环境说明.docx')
            return ChenResponse(status=200, code=200, message="文档生成成功！")
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

    # 静态软件项
    @route.get('/create/static_soft', url_name='create-static_soft')
    def create_static_soft(self, id: int):
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '静态软件项.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '静态软件项')
        context = {
            "replace": replace,  # 指定是否由数据库文档片段进行生成
            "user_content": frag and rich_text_list
        }
        return create_dg_docx("静态软件项.docx", context, id)

    # 静态硬件和固件项
    @route.get('/create/static_hard', url_name='create-static_hard')
    def create_static_hard(self, id: int):
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '静态硬件和固件项.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '静态硬件和固件项')
        context = {
            "replace": replace,  # 指定是否由数据库文档片段进行生成
            "user_content": frag and rich_text_list
        }
        return create_dg_docx("静态硬件和固件项.docx", context, id)

    # 动态测评环境说明
    @route.get('/create/dynamic_env', url_name='create-dynamic_env')
    def create_dynamic_env(self, id: int):
        project_obj: Project = get_object_or_404(Project, id=id)
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '动态测试环境说明.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '动态测试环境说明')
        context = {
            'project_name': project_obj.name,
            "replace": replace,
            "user_content": frag and rich_text_list
        }
        doc.render(context, autoescape=True)
        try:
            doc.save(Path.cwd() / "media" / project_path(id) / "output_dir" / '动态测试环境说明.docx')
            return ChenResponse(status=200, code=200, message="文档生成成功！")
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))

    # 动态软件项
    @route.get('/create/dynamic_soft', url_name='create-dynamic_soft')
    def create_dynamic_soft(self, id: int):
        project_obj: Project = get_object_or_404(Project, id=id)
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '动态软件项.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '动态软件项')
        context = {
            'project_name': project_obj.name,
            "replace": replace,
            "user_content": frag and rich_text_list
        }
        return create_dg_docx("动态软件项.docx", context, id)

    # 动态软件项
    @route.get('/create/dynamic_hard', url_name='create-dynamic_hard')
    def create_dynamic_hard(self, id: int):
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '动态硬件和固件项.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '动态硬件和固件项')
        context = {
            "replace": replace,
            "user_content": frag and rich_text_list
        }
        return create_dg_docx("动态硬件和固件项.docx", context, id)

    # 测试数据
    @route.get('/create/test_data', url_name='create-test_data')
    def create_test_data(self, id: int):
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '测评数据.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '测评数据')
        context = {
            "replace": replace,
            "user_content": frag and rich_text_list
        }
        return create_dg_docx("测评数据.docx", context, id)

    # 环境差异性分析
    @route.get('/create/env_diff', url_name='create-env_diff')
    def create_env_diff(self, id: int):
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '环境差异性分析.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '环境差异性分析')
        context = {
            "replace": replace,
            "user_content": frag and rich_text_list
        }
        return create_dg_docx("环境差异性分析.docx", context, id)

    # 生成被测软件-基本信息
    @route.get('/create/baseInformation', url_name='create-baseInformation')
    def create_information(self, id: int):
        project_qs = get_object_or_404(Project, id=id)
        security = get_str_dict(project_qs.security_level, 'security_level')
        languages = get_list_dict('language', project_qs.language)
        runtime = get_str_dict(project_qs.runtime, 'runtime')
        devplant = get_str_dict(project_qs.devplant, 'devplant')
        language_list = []
        for language in languages:
            language_list.append(language.get('ident_version'))
        # 版本先找第一轮
        project_round = project_qs.pField.filter(key=0).first()
        first_round_SO = project_round.rdField.filter(type='SO').first()
        if not first_round_SO:
            return ChenResponse(code=400, status=400, message='您还未创建轮次，请进入工作区创建')
        version = first_round_SO.version
        line_count = int(first_round_SO.total_lines)
        dev_unit = project_qs.dev_unit
        # 渲染上下文
        context = {
            'project_name': project_qs.name,
            'is_JD': True if project_qs.report_type == '9' else False,
            'security_level': security,
            'language': "\a".join(language_list),
            'version': version,
            'line_count': line_count,
            'effective_line': int(first_round_SO.effective_lines),
            'recv_date': project_qs.beginTime.strftime("%Y-%m-%d"),
            'dev_unit': dev_unit,
            'soft_type': project_qs.get_soft_type_display(),
            'runtime': runtime,
            'devplant': devplant
        }
        return create_dg_docx('被测软件基本信息.docx', context, id)

    # 生成测试级别和测试类型
    @route.get('/create/levelAndType', url_name='create-levelAndType')
    def create_levelAndType(self, id: int):
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '测试级别和测试类型.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '测试级别和测试类型')
        if replace:
            context = {
                "replace": replace,
                "user_content": frag and rich_text_list
            }
        else:
            # 如果没有片段替换，则利用数据生成信息
            project_qs = get_object_or_404(Project, id=id)
            # 获取所有已录入测试类型
            test_types = project_qs.ptField.values("testType").distinct()
            # 通过测试类型查询字典中的中文
            type_name_list = list(
                map(lambda qs_item: get_str_dict(qs_item['testType'], 'testType'), test_types))
            # 定义测试类型一览的顺序，注意word里面也要一样
            word_types = ['文档审查', '静态分析', '代码审查', '逻辑测试', '功能测试', '性能测试', '边界测试',
                          '恢复性测试', '安装性测试', '数据处理测试', '余量测试', '强度测试', '接口测试',
                          '人机交互界面测试', '兼容性测试']
            type_index = []
            for index, test_type in enumerate(word_types):
                for exist_type in type_name_list:
                    if exist_type == test_type:
                        type_index.append(str(index))
            context = {
                "security_level": get_str_dict(project_qs.security_level, 'security_level'),
                "testTypes": "、".join(type_name_list),
                "project_name": project_qs.name,
                "type_index": type_index
            }
        return create_dg_docx("测试级别和测试类型.docx", context, id)

    # 生成测试策略
    @route.get('/create/strategy', url_name='create-strategy')
    def create_strategy(self, id: int):
        input_path = Path.cwd() / 'media' / project_path(id) / 'form_template' / 'dg' / '测试策略.docx'
        doc = DocxTemplate(input_path)
        replace, frag, rich_text_list = self._generate_frag(id, doc, '测试策略')
        if replace:
            context = {
                "replace": replace,
                "user_content": frag and rich_text_list
            }
        else:
            # 如果没有片段替换，则利用数据生成信息
            project_qs = get_object_or_404(Project, id=id)
            # 根据关键等级检查是否有代码审查
            security = project_qs.security_level
            isDmsc = True if int(security) <= 2 else False
            # 获取第一轮所有测试项QuerySet
            project_round_one = project_qs.pField.filter(key=0).first()
            testDemand_qs = project_round_one.rtField.all()
            # grouped_data的键是测试类型名称，值为测试项名称数组
            grouped_data = {}
            for item in testDemand_qs:
                grouped_data.setdefault(get_str_dict(item.testType, 'testType'), []).append(item.name)
            # 获取当前测试项的测试类型
            test_types = testDemand_qs.values("testType").distinct()
            type_name_list = list(
                map(lambda qs_item: get_str_dict(qs_item['testType'], 'testType'), test_types))
            context = {
                "project_name": project_qs.name,
                # 查询关键等级-类似“关键”输出
                "security_level_str": get_str_abbr(security, 'security_level'),
                "isDmsc": isDmsc,
                "test_types": type_name_list,
                "grouped_data": grouped_data,
            }
        return create_dg_docx("测试策略.docx", context, id)

    # 生成-测试内容充分性及测试方法有效性
    @route.get('/create/adequacy_effectiveness', url_name='create-adequacy_effectiveness')
    def create_adequacy_effectiveness(self, id: int):
        project_qs = get_object_or_404(Project, id=id)
        # 统计测试种类数量-只统计第一轮测试
        project_round_one = project_qs.pField.filter(key=0).first()
        if not project_round_one:
            return ChenResponse(status=400, code=400, message="未找到首轮测试信息!")
        # 通过字典获取-测试方法
        type_dict = {}  # key为测试类型，value为数量
        testDemands = project_round_one.rtField.all()
        for testDemand in testDemands:
            # 获取每个测试项测试类型
            test_type = get_list_dict('testType', [testDemand.testType])[0].get('ident_version')
            # 如果字典没有该key，则创建并value=1
            if not test_type in type_dict:
                type_dict[test_type] = 1
            else:
                type_dict[test_type] += 1
        length = len(type_dict)
        type_str_list = []
        for key, value in type_dict.items():
            type_str_list.append(f"{key}{value}项")
        context = {
            'project_name': project_qs.name,
            'test_item_count': testDemands.count(),
            'length': length,
            'type_str': "、".join(type_str_list),
        }
        return create_dg_docx('测试内容充分性及测试方法有效性分析.docx', context, id)

    # 生成-测评项目组组成和分工
    @route.get('/create/group', url_name='create_group')
    def create_group(self, id: int):
        project_qs = get_object_or_404(Project, id=id)
        context = {
            'duty_person': project_qs.duty_person,
            'member_str': "、".join(project_qs.member),
            'quality_person': project_qs.quality_person,
            'vise_person': project_qs.vise_person,
            'config_person': project_qs.config_person,
            'dev_unit': project_qs.dev_unit,
        }
        return create_dg_docx('测评组织及任务分工.docx', context, id)

    # 生成-测评条件保障
    @route.get('/create/guarantee', url_name='create-guarantee')
    def create_guarantee(self, id: int):
        project_qs = get_object_or_404(Project, id=id)
        context = {
            'project': project_qs
        }
        return create_dg_docx('测评条件保障.docx', context, id)

    # 生成-缩略语
    @route.get('/create/abbreviation', url_name='create-abbreviation')
    def create_abbreviation(self, id: int):
        project_qs = get_object_or_404(Project, id=id)
        abbreviations = []
        for abbr in project_qs.abbreviation:
            abbr_dict = {'title': abbr, 'des': Abbreviation.objects.filter(title=abbr).first().des}
            abbreviations.append(abbr_dict)
        context = {
            'abbreviations': abbreviations
        }
        return create_dg_docx('缩略语.docx', context, id)

    # 生成研制总要求-测试项追踪关系表
    @route.get('/create/yzComparison', url_name='create-yzComparison')
    def create_yzComparison(self, id: int):
        """目前追踪需求项的章节号是硬编码，按6.2章节起步，6.2.1~x.x.x依次排序"""
        # 规定测试项的章节号开头
        test_item_prefix = '6.2'
        # 计算有多少种testType - '文档审查'/'功能测试' ->
        # 形成一个数组['1','2','3','4','9']后面用来判断测试项的章节号
        project_qs = get_object_or_404(Project, id=id)
        design_list = []  # 先按照design的思路进行追踪
        # 查询第一轮次
        project_round_one = project_qs.pField.filter(key=0).first()
        testType_list, last_chapter_items = create_csx_chapter_dict(project_round_one)
        # 找出第一轮的研总
        yz_dut = project_round_one.rdField.filter(type='YZ').first()
        if yz_dut:
            # 查询出验证所有design
            yz_designs = yz_dut.rsField.all()
            # 遍历所有研总的design
            for design in yz_designs:
                design_dict = {'name': design.name, 'chapter': design.chapter, 'test_demand': []}
                # 获取一个design的所有测试项
                test_items = design.dtField.all()
                # 连接两个QuerySet，默认去重
                test_items = test_items.union(design.odField.all())
                for test_item in test_items:
                    reveal_ident = "_".join(
                        ["XQ", get_testType(test_item.testType, "testType"), test_item.ident])
                    # 查字典方式确认章节号最后一位
                    test_item_last_chapter = last_chapter_items[test_item.testType].index(
                        test_item.key) + 1
                    test_chapter = ".".join(
                        [test_item_prefix, str(testType_list.index(test_item.testType) + 1),
                         str(test_item_last_chapter)])
                    test_item_dict = {'name': test_item.name, 'chapter': test_chapter,
                                      'ident': reveal_ident}
                    design_dict['test_demand'].append(test_item_dict)
                design_list.append(design_dict)
        try:
            design_list = sorted(design_list, key=chapter_key)
        except Exception as e:
            print("研总的追踪排序报错，错误原因：", e)
        context = {
            'design_list': design_list
        }
        return create_dg_docx('研制总要求追踪表.docx', context, id)

    # 生成需求规格说明-测试项追踪关系表
    @route.get('/create/xqComparison', url_name='create-xqComparison')
    def create_xqComparison(self, id: int):
        project_qs = get_object_or_404(Project, id=id)
        test_item_prefix = '6.2'
        design_list = []
        project_round_one = project_qs.pField.filter(key=0).first()
        if project_round_one:
            testType_list, last_chapter_items = create_csx_chapter_dict(project_round_one)
            # 找出第一轮的被测件为'XQ'
            xq_dut = project_round_one.rdField.filter(type='XQ').first()
            # 找出第一轮被测件为'SO'，其中的测试项
            so_dut = project_round_one.rdField.filter(type='SO').first()
            if so_dut:
                so_designs = so_dut.rsField.all()
                for design in so_designs:
                    design_dict = {'name': "/", 'chapter': "/", 'test_demand': []}
                    # 获取一个design的所有测试项
                    test_items = []
                    test_items.extend(design.dtField.all())
                    test_items.extend(design.odField.all())

                    for test_item in test_items:
                        # 只对文档审查、静态分析、代码走查、代码审查进行处理
                        if test_item.testType in ['8', '15', '3', '2']:
                            reveal_ident = "_".join(
                                ["XQ", get_testType(test_item.testType, "testType"), test_item.ident])
                            # 查字典方式确认章节号最后一位
                            test_item_last_chapter = last_chapter_items[test_item.testType].index(
                                test_item.key) + 1
                            test_chapter = ".".join(
                                [test_item_prefix, str(testType_list.index(test_item.testType) + 1),
                                 str(test_item_last_chapter)])
                            test_item_dict = {'name': test_item.name, 'chapter': test_chapter,
                                              'ident': reveal_ident}
                            design_dict['test_demand'].append(test_item_dict)
                    design_list.append(design_dict)

            if xq_dut:
                xq_designs = xq_dut.rsField.all()
                for design in xq_designs:
                    design_dict = {'name': design.name, 'chapter': design.chapter, 'test_demand': []}
                    # 获取一个design的所有测试项
                    test_items = []
                    test_items.extend(design.dtField.all())
                    test_items.extend(design.odField.all())

                    for test_item in test_items:
                        reveal_ident = "_".join(
                            ["XQ", get_testType(test_item.testType, "testType"), test_item.ident])
                        # 查字典方式确认章节号最后一位
                        test_item_last_chapter = last_chapter_items[test_item.testType].index(
                            test_item.key) + 1
                        test_chapter = ".".join(
                            [test_item_prefix, str(testType_list.index(test_item.testType) + 1),
                             str(test_item_last_chapter)])
                        test_item_dict = {'name': test_item.name, 'chapter': test_chapter,
                                          'ident': reveal_ident}
                        design_dict['test_demand'].append(test_item_dict)

                    design_list.append(design_dict)
            # 根据design的chapter排序-为防止报错崩溃使用try-但难排查
            try:
                design_list = sorted(design_list, key=chapter_key)
            except Exception as e:
                print("追踪排序报错，错误原因：", e)
            context = {
                'design_list': design_list
            }
            return create_dg_docx('需求规格说明追踪表.docx', context, id)
        raise HttpError(400, "生成需求追踪表出错")

    # 生成测试项-需求规格说明关系表【反向】
    @route.get('/create/fanXqComparison', url_name='create-fanXqComparison')
    def create_fanXqComparison(self, id: int):
        project_qs = get_object_or_404(Project, id=id)
        test_item_prefix = '6.2'
        # 取出第一轮所有测试项的章节处理列表和字典
        project_round_one = project_qs.pField.filter(key=0).first()
        testType_list, last_chapter_items = create_csx_chapter_dict(project_round_one)
        # 查询第一轮所有测试项
        test_items = []
        test_items.extend(project_round_one.rtField.all())
        # 最后渲染列表
        items_list = []
        for test_item in test_items:
            # 第二个处理被测件为"XQ"，第二个处理被测件为'SO'，并且为测试项testType为['8', '15', '3', '2']的
            if test_item.dut.type == 'XQ' or (
                    test_item.dut.type == 'SO' and test_item.testType in ['8', '15', '3',
                                                                          '2']):
                reveal_ident = "_".join(
                    ["XQ", get_testType(test_item.testType, "testType"), test_item.ident])
                # 查字典方式确认章节号最后一位
                test_item_last_chapter = last_chapter_items[test_item.testType].index(test_item.key) + 1
                test_chapter = ".".join([test_item_prefix, str(testType_list.index(test_item.testType) + 1),
                                         str(test_item_last_chapter)])
                # 如果是SO里面的
                if test_item.testType in ['8', '15', '3', '2'] and test_item.dut.type == 'SO':
                    test_item_dict = {'name': test_item.name, 'chapter': test_chapter, 'ident': reveal_ident,
                                      'design': {
                                          'name': "/", 'chapter': "/"
                                      }}
                else:
                    test_item_dict = {'name': test_item.name, 'chapter': test_chapter, 'ident': reveal_ident,
                                      'design': {
                                          'name': test_item.design.name, 'chapter': test_item.design.chapter
                                      }}
                items_list.append(test_item_dict)
        context = {
            'items_list': items_list,
        }
        return create_dg_docx('反向需求规格追踪表.docx', context, id)

    # 生成代码质量度量分析表
    @route.get('/create/codeQuality', url_name='create-codeQuality')
    def create_codeQuality(self, id: int):
        project_qs = get_object_or_404(Project, id=id)
        project_round_one = project_qs.pField.filter(key=0).first()
        context = {}
        context.update({'project_name': project_qs.name})
        if project_round_one:
            source_dut: Dut = project_round_one.rdField.filter(type='SO').first()  # type:ignore
            if source_dut:
                context.update({'version': source_dut.version})  # type:ignore
                context.update({'size': int(source_dut.total_lines)})
                context.update({'total_code_line': int(source_dut.effective_lines)})
                context.update({'comment_line': int(source_dut.comment_lines)})
                comment_ratio = int(source_dut.comment_lines) / int(source_dut.total_lines)
                context.update({
                    'comment_ratio': f"{comment_ratio * 100:.2f}%",
                    'comment_ratio_right': '满足' if comment_ratio >= 0.2 else '不满足'
                })
                # 如果已经有metrics
                if hasattr(source_dut, 'metrics'):
                    context.update({
                        'black_line': source_dut.metrics.total_blanks,
                        'function_count': source_dut.metrics.function_count,
                        'avg_function_lines': source_dut.metrics.avg_function_lines,
                        'avg_function_lines_right': '满足' if source_dut.metrics.avg_function_lines <= 200 else '不满足',
                        'avg_fan_out': source_dut.metrics.avg_fan_out,
                        'avg_fan_out_right': '满足' if source_dut.metrics.avg_fan_out <= 7 else '不满足',
                        'avg_cyclomatic': source_dut.metrics.avg_cyclomatic,
                        'avg_cyclomatic_right': '满足' if source_dut.metrics.avg_cyclomatic <= 10 else '不满足',
                        'max_cyclomatic': source_dut.metrics.max_cyclomatic,
                        'max_cyclomatic_right': '满足' if source_dut.metrics.max_cyclomatic <= 80 else '不满足',
                        'high_cyclomatic_ratio': source_dut.metrics.high_cyclomatic_ratio,
                        'high_cyclomatic_ratio_right': '满足' if source_dut.metrics.high_cyclomatic_ratio <= 0.2 else '不满足',
                    })
            else:
                return ChenResponse(message='未找到源代码被测件', code=400)
        return create_dg_docx('代码质量度量分析表.docx', context, id)

# 工具方法-给sorted排序使用-知识点：python里面可以元组排序
def chapter_key(item):
    big_num = [5000, 5000, 5000, 5000]
    if "." in item['chapter']:
        # 如果是有章节号的则排序
        return [int(part) for part in item['chapter'].split(".")]
    if item['test_demand'][0]['name'] in ['文档审查', '静态分析', '代码审查', '代码走查']:
        return [0, 0, 0, 0]
    return big_num
